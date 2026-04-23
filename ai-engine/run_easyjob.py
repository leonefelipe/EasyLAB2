"""
run_easyjob.py — Runner CLI do Método EasyJob

Uso:
    python run_easyjob.py init cliente_x.json
    python run_easyjob.py run cliente_x.json

Saídas em saida/<cliente>/:
    01_CV_Otimizado.docx / .md
    02_Roteiro_LinkedIn.md / .json
    03_Relatorio_Premium.docx / .md
    _diagnostico.json
"""

from __future__ import annotations

import argparse
import json
import os
import sys
from pathlib import Path

from easyjob_engine import (
    DadosCandidato,
    Entregaveis,
    LLMClient,
    gerar_entregaveis_sync,
)


TEMPLATE_INPUT = {
    "nome": "Nome Completo do Candidato",
    "contato": {
        "email": "email@dominio.com",
        "telefone": "+55 11 90000-0000",
        "cidade": "São Paulo, SP",
        "linkedin": "linkedin.com/in/usuario",
    },
    "titulo_atual": "Cargo atual | Setor | Foco técnico",
    "resumo_atual": "Cole aqui o resumo/objetivo atualmente escrito no CV do cliente.",
    "experiencias": [
        {
            "cargo": "Cargo",
            "empresa": "Empresa",
            "local": "Cidade, UF",
            "periodo": "MM/AAAA – MM/AAAA",
            "anos": 0.0,
            "bullets": [
                "Cole cada bullet atual em uma string separada, na ordem original do CV",
            ],
        }
    ],
    "formacao": [
        {
            "curso": "Graduação / Pós",
            "instituicao": "Instituição",
            "periodo": "AAAA – AAAA",
            "detalhe": "",
        }
    ],
    "certificacoes": [],
    "idiomas": [{"idioma": "Inglês", "nivel": "Avançado"}],
    "habilidades_declaradas": [],
    "vaga_alvo": {
        "titulo": "Cargo alvo",
        "empresa": "Empresa alvo (opcional)",
        "senioridade": "Sênior / Head / Diretor",
        "anos_requeridos": 5,
        "descricao": "Cole aqui a descrição da vaga-alvo (job description).",
    },
    "objetivo_carreira": "Em uma frase, o alvo estratégico de 12 meses.",
}


# ═══════════════════════════════════════════════════════════════════════════════
# EXPORTADORES DOCX
# ═══════════════════════════════════════════════════════════════════════════════

def exportar_cv_docx(cv_estrutura: dict, destino: Path) -> None:
    """CV em DOCX com formatação ATS-aligned (uma coluna, headings padronizados)."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.3)
        section.bottom_margin = Cm(1.3)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    cab = cv_estrutura["cabecalho"]

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(cab["nome"])
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    run = p.add_run(cab["titulo_profissional"])
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(cab["linha_contato"])

    def h2(texto: str) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(texto.upper())
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x1B, 0x2F, 0x4A)  # navy (brand book Leone Berto)

    h2("Resumo Executivo")
    doc.add_paragraph(cv_estrutura["resumo_executivo"])

    h2("Experiência Profissional")
    for e in cv_estrutura.get("experiencias", []):
        p = doc.add_paragraph()
        r = p.add_run(f"{e['cargo']} — {e['empresa']}")
        r.bold = True
        meta_partes = [x for x in [e.get("local", ""), e.get("periodo", "")] if x]
        if meta_partes:
            p2 = doc.add_paragraph()
            r2 = p2.add_run(" · ".join(meta_partes))
            r2.italic = True
        if e.get("contexto"):
            doc.add_paragraph(e["contexto"])
        for b in e.get("bullets", []):
            doc.add_paragraph(b["texto"], style="List Bullet")

    if cv_estrutura.get("formacao"):
        h2("Formação")
        for f in cv_estrutura["formacao"]:
            linha = f"{f.get('curso','')} — {f.get('instituicao','')}"
            if f.get("periodo"):
                linha += f" ({f['periodo']})"
            doc.add_paragraph(linha, style="List Bullet")

    if cv_estrutura.get("certificacoes"):
        h2("Certificações")
        for c in cv_estrutura["certificacoes"]:
            doc.add_paragraph(c, style="List Bullet")

    if cv_estrutura.get("idiomas"):
        h2("Idiomas")
        for i in cv_estrutura["idiomas"]:
            doc.add_paragraph(f"{i['idioma']}: {i['nivel']}", style="List Bullet")

    comp = cv_estrutura.get("competencias", {})
    if any(comp.values()):
        h2("Competências")
        mapa = [
            ("tecnicas", "Técnicas"),
            ("ferramentas", "Ferramentas"),
            ("negocios", "Negócios"),
            ("lideranca", "Liderança"),
        ]
        for chave, rotulo in mapa:
            if comp.get(chave):
                p = doc.add_paragraph()
                r = p.add_run(f"{rotulo}: ")
                r.bold = True
                p.add_run(", ".join(comp[chave]))

    doc.save(str(destino))


def exportar_relatorio_premium_docx(markdown: str, destino: Path, nome_cliente: str) -> None:
    """Relatório Premium em DOCX com capa (identidade Leone Berto: navy + dourado)."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    navy = RGBColor(0x1B, 0x2F, 0x4A)
    gold = RGBColor(0xC8, 0xA1, 0x5E)

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MÉTODO EASYJOB")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = navy

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Relatório Premium de Reposicionamento Executivo")
    r.font.size = Pt(14)
    r.font.color.rgb = gold
    r.italic = True

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Preparado para: {nome_cliente}")
    r.font.size = Pt(12)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Felipe Leone | Estratégia de Carreira e Posicionamento Profissional")
    r.font.size = Pt(10)
    r.italic = True

    doc.add_page_break()

    for linha in markdown.splitlines():
        linha = linha.rstrip()
        if not linha:
            doc.add_paragraph()
        elif linha.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(linha[2:])
            r.bold = True
            r.font.size = Pt(16)
            r.font.color.rgb = navy
        elif linha.startswith("## "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            r = p.add_run(linha[3:])
            r.bold = True
            r.font.size = Pt(13)
            r.font.color.rgb = navy
        elif linha.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            r = p.add_run(linha[4:])
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = gold
        elif linha.startswith("- "):
            doc.add_paragraph(linha[2:], style="List Bullet")
        elif linha.startswith("---"):
            doc.add_paragraph("_" * 60)
        else:
            _adicionar_paragrafo_com_negrito(doc, linha)

    doc.save(str(destino))


def _adicionar_paragrafo_com_negrito(doc, texto: str) -> None:
    import re as _re

    p = doc.add_paragraph()
    partes = _re.split(r"(\*\*[^*]+\*\*)", texto)
    for parte in partes:
        if parte.startswith("**") and parte.endswith("**"):
            r = p.add_run(parte[2:-2])
            r.bold = True
        else:
            p.add_run(parte)


def renderizar_roteiro_linkedin_markdown(r: dict) -> str:
    linhas = [
        "# Roteiro Operacional de Otimização de LinkedIn",
        "## Método EasyJob por Felipe Leone",
        "",
        "## 1. Headline",
        f"**Atual:** {r['headline'].get('atual','(vazio)')}",
        "",
        f"**Recomendado:** {r['headline']['recomendado']}",
        "",
        f"> *{r['headline']['racional']}*",
        "",
        "## 2. Seção Sobre",
        f"**Resumo do atual:** {r['sobre'].get('atual_resumo','(vazio)')}",
        "",
        "**Texto recomendado — pronto para colar:**",
        "",
        "```",
        r["sobre"]["recomendado_completo"],
        "```",
        "",
        f"> *{r['sobre']['racional']}*",
        "",
        "## 3. Experiência",
    ]
    for e in r.get("experiencia", []):
        linhas.append(f"\n### {e['cargo']} — {e['empresa']}")
        linhas.append("Alterações prioritárias:")
        for a in e.get("alteracoes_prioritarias", []):
            linhas.append(f"- {a}")
        if e.get("bullet_modelo"):
            linhas.append(f"\n**Bullet modelo:** {e['bullet_modelo']}")

    linhas += [
        "",
        "## 4. Competências",
        f"**Fixar no topo (3):** {', '.join(r.get('competencias_fixadas_top3', []))}",
        "",
        f"**Adicionar:** {', '.join(r.get('competencias_adicionar', []))}",
        "",
        f"**Remover:** {', '.join(r.get('competencias_remover', []))}",
        "",
        "## 5. Seção Em Destaque (Featured)",
    ]
    for f in r.get("featured_sugestoes", []):
        linhas.append(f"- {f}")

    linhas += [
        "",
        "## 6. Recomendações",
        f"Meta mínima: {r.get('recomendacoes', {}).get('meta', 5)} recomendações escritas.",
        "",
        "Alvos sugeridos:",
    ]
    for a in r.get("recomendacoes", {}).get("alvos_sugeridos", []):
        linhas.append(f"- {a}")

    linhas += ["", "## 7. Plano de Conteúdo (8 semanas)"]
    for sem in r.get("plano_conteudo_8_semanas", []):
        linhas.append(f"\n**Semana {sem.get('semana','?')}**")
        for k in ("post_1", "post_2", "post_3"):
            if sem.get(k):
                linhas.append(f"- {sem[k]}")

    hig = r.get("higiene_algoritmica", {})
    linhas += [
        "",
        "## 8. Higiene Algorítmica",
        f"**Completude atual do perfil:** {hig.get('completude_atual_pct', 0)}%",
        "",
        "**Ações imediatas:**",
    ]
    for a in hig.get("acoes_imediatas", []):
        linhas.append(f"- {a}")
    if hig.get("indicadores_ssi_foco"):
        linhas.append("")
        linhas.append(f"**Dimensões de SSI a priorizar:** {', '.join(hig['indicadores_ssi_foco'])}")

    linhas += ["", "## 9. Checklist de Execução"]
    for c in r.get("checklist_execucao", []):
        linhas.append(f"- [ ] {c}")

    return "\n".join(linhas)


# ═══════════════════════════════════════════════════════════════════════════════
# COMANDOS CLI
# ═══════════════════════════════════════════════════════════════════════════════

def cmd_init(caminho: str) -> int:
    destino = Path(caminho)
    if destino.exists():
        print(f"[!] {destino} já existe. Aborte ou renomeie.", file=sys.stderr)
        return 1
    destino.write_text(
        json.dumps(TEMPLATE_INPUT, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(f"[ok] Template gerado em {destino}")
    print("     Preencha os campos e rode: python run_easyjob.py run " + caminho)
    return 0


def cmd_run(caminho: str) -> int:
    if not os.environ.get("OPENAI_API_KEY"):
        print("[!] OPENAI_API_KEY não definida no ambiente.", file=sys.stderr)
        return 2

    origem = Path(caminho)
    if not origem.exists():
        print(f"[!] Arquivo {origem} não encontrado.", file=sys.stderr)
        return 1

    dados = json.loads(origem.read_text(encoding="utf-8"))
    candidato = DadosCandidato(**dados)

    slug = _slug(candidato.nome) or origem.stem
    saida = Path("saida") / slug
    saida.mkdir(parents=True, exist_ok=True)

    print(f"[..] Gerando entregáveis para: {candidato.nome}")
    print(f"     Modelo OpenAI: {os.environ.get('EASYJOB_MODEL', 'gpt-4o')}")
    print(f"     Destino: {saida.resolve()}")

    llm = LLMClient()
    resultado: Entregaveis = gerar_entregaveis_sync(candidato, llm)

    (saida / "01_CV_Otimizado.md").write_text(
        resultado.cv_otimizado_markdown, encoding="utf-8"
    )
    exportar_cv_docx(
        resultado.cv_otimizado_estrutura,
        saida / "01_CV_Otimizado.docx",
    )

    (saida / "02_Roteiro_LinkedIn.json").write_text(
        json.dumps(resultado.roteiro_linkedin, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    (saida / "02_Roteiro_LinkedIn.md").write_text(
        renderizar_roteiro_linkedin_markdown(resultado.roteiro_linkedin),
        encoding="utf-8",
    )

    (saida / "03_Relatorio_Premium.md").write_text(
        resultado.relatorio_premium_markdown, encoding="utf-8"
    )
    exportar_relatorio_premium_docx(
        resultado.relatorio_premium_markdown,
        saida / "03_Relatorio_Premium.docx",
        candidato.nome,
    )

    (saida / "_diagnostico.json").write_text(
        json.dumps(resultado.diagnostico, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    print("[ok] Entregáveis gerados:")
    for arquivo in sorted(saida.iterdir()):
        print(f"     - {arquivo.name}")
    print(f"\n[score] ATS atual diagnosticado: {resultado.diagnostico['ats_score_atual']}/100")
    return 0


def _slug(texto: str) -> str:
    import re as _re
    import unicodedata as _u

    t = _u.normalize("NFD", texto.lower())
    t = "".join(c for c in t if _u.category(c) != "Mn")
    t = _re.sub(r"[^a-z0-9]+", "_", t).strip("_")
    return t[:40]


def main() -> int:
    parser = argparse.ArgumentParser(
        prog="run_easyjob",
        description="Método EasyJob — runner CLI dos 3 entregáveis (OpenAI)",
    )
    sub = parser.add_subparsers(dest="cmd", required=True)

    p_init = sub.add_parser("init", help="gerar template JSON de entrada")
    p_init.add_argument("arquivo", help="caminho do JSON a criar")

    p_run = sub.add_parser("run", help="rodar o motor sobre um JSON preenchido")
    p_run.add_argument("arquivo", help="caminho do JSON do candidato")

    args = parser.parse_args()
    if args.cmd == "init":
        return cmd_init(args.arquivo)
    if args.cmd == "run":
        return cmd_run(args.arquivo)
    return 1


if __name__ == "__main__":
    sys.exit(main())
